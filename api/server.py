import os
import httpx
import json
import io
import re
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import List, Optional
from celery.result import AsyncResult
from worker.celery_app import celery_app
from worker.tasks import run_research
from sqlalchemy import create_engine, text
from dotenv import load_dotenv
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm
from reportlab.lib import colors

load_dotenv()

app = FastAPI(title="Research Agent API", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

security = HTTPBearer()
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")


def get_engine():
    return create_engine(os.getenv("DATABASE_URL"), connect_args={"sslmode": "require"})


async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    async with httpx.AsyncClient() as client:
        response = await client.get(
            f"{SUPABASE_URL}/auth/v1/user",
            headers={"Authorization": f"Bearer {token}", "apikey": SUPABASE_ANON_KEY}
        )
    if response.status_code != 200:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    return response.json()


class ResearchRequest(BaseModel):
    topic: str
    pdf_paths: Optional[List[str]] = []


@app.get("/")
def home():
    return {"status": "Research Agent running"}


@app.post("/upload-pdfs")
async def upload_pdfs(
    files: List[UploadFile] = File(...),
    user=Depends(get_current_user)
):
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Maximum 20 PDFs allowed")

    from supabase import create_client
    sb = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)

    uploaded_paths = []
    for file in files:
        if not file.filename.endswith(".pdf"):
            raise HTTPException(status_code=400, detail=f"{file.filename} is not a PDF")
        content = await file.read()
        path = f"{user['id']}/{file.filename}"
        sb.storage.from_("research-pdfs").upload(
            path, content, {"content-type": "application/pdf", "upsert": "true"}
        )
        uploaded_paths.append(path)

    return {"uploaded": uploaded_paths}


@app.post("/research")
async def research(request: ResearchRequest, user=Depends(get_current_user)):
    job = run_research.delay(request.topic, user["id"], request.pdf_paths)
    return {"job_id": job.id, "status": "Research started"}


@app.get("/status/{job_id}")
async def status(job_id: str, user=Depends(get_current_user)):
    engine = get_engine()
    with engine.connect() as conn:
        row = conn.execute(
            text("SELECT status, result FROM research_jobs WHERE job_id = :job_id AND user_id = :uid"),
            {"job_id": job_id, "uid": user["id"]}
        ).fetchone()

    if not row:
        celery_job = AsyncResult(job_id, app=celery_app)
        return {"job_id": job_id, "status": celery_job.status, "result": None}

    result = json.loads(row.result) if row.result else None
    celery_status = "SUCCESS" if row.status == "done" else "PENDING" if row.status == "running" else "FAILURE"
    return {"job_id": job_id, "status": celery_status, "result": result}


@app.get("/jobs")
async def get_jobs(user=Depends(get_current_user)):
    engine = get_engine()
    with engine.connect() as conn:
        result = conn.execute(
            text("SELECT job_id, topic, status, created_at FROM research_jobs WHERE user_id = :uid ORDER BY created_at DESC"),
            {"uid": user["id"]}
        )
        rows = result.fetchall()
    return {"jobs": [dict(r._mapping) for r in rows]}


def clean_text(text):
    return re.sub(r'[^\x00-\x7F]+', '', text or "")


@app.get("/download/{job_id}/{format}")
async def download(job_id: str, format: str, user=Depends(get_current_user)):
    engine = get_engine()
    with engine.connect() as conn:
        row = conn.execute(
            text("SELECT topic, result FROM research_jobs WHERE job_id = :job_id AND user_id = :uid"),
            {"job_id": job_id, "uid": user["id"]}
        ).fetchone()

    if not row or not row.result:
        raise HTTPException(status_code=404, detail="Job not found or not completed")

    data = json.loads(row.result)
    topic = row.topic

    if format == "docx":
        doc = Document()
        doc.add_heading(f"Research Report: {topic}", 0)

        doc.add_heading("Paper Summaries", level=1)
        for s in data.get("summaries", []):
            doc.add_heading(s["title"], level=2)
            doc.add_paragraph(s["summary"])
            doc.add_paragraph("")

        doc.add_heading("Research Insights", level=1)
        doc.add_paragraph(data.get("insights", ""))

        doc.add_heading("Critical Review", level=1)
        doc.add_paragraph(data.get("critique", ""))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"research_{topic[:30].replace(' ', '_')}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    elif format == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=18, spaceAfter=20)
        h1_style = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=14, spaceAfter=10, spaceBefore=16)
        h2_style = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12, spaceAfter=8, spaceBefore=12)
        body_style = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, spaceAfter=6, leading=16)

        story = []
        story.append(Paragraph(clean_text(f"Research Report: {topic}"), title_style))
        story.append(Spacer(1, 0.3*cm))

        story.append(Paragraph("Paper Summaries", h1_style))
        for s in data.get("summaries", []):
            story.append(Paragraph(clean_text(s["title"]), h2_style))
            for line in clean_text(s["summary"]).split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 0.3*cm))

        story.append(Paragraph("Research Insights", h1_style))
        for line in clean_text(data.get("insights", "")).split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), body_style))

        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph("Critical Review", h1_style))
        for line in clean_text(data.get("critique", "")).split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), body_style))

        doc.build(story)
        buf.seek(0)
        filename = f"research_{topic[:30].replace(' ', '_')}.pdf"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    raise HTTPException(status_code=400, detail="Format must be pdf or docx")
